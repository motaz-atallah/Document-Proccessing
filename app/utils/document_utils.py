import os
import fitz  # PyMuPDF
from docx import Document

def extract_text_from_pdf(pdf_file):
    """Extract text from a PDF file."""
    text = ""
    doc = fitz.open(pdf_file)
    for page_num in range(doc.page_count):
        page = doc[page_num]
        text += page.get_text("text")
    return text

def extract_text_from_word(word_file):
    """Extract text from a Word file."""
    doc = Document(word_file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_text_from_word(word_file):
    """Extract text from a Word file."""
    doc = Document(word_file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def process_document(file_path):
    """Process the document and extract information based on the file type."""
    ext = os.path.splitext(file_path)[1]

    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_word(file_path)
    else:
        raise ValueError("Unsupported file format.")
